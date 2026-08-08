from __future__ import annotations

import re
from collections.abc import Callable
from io import BytesIO
from uuid import uuid4

import fitz
from docx import Document

from app.domain.source import SourceSpan


class DocumentParseError(ValueError):
    """Raised when an allowed document cannot be parsed safely."""


def _span_id() -> str:
    return f"span_{uuid4().hex}"


def _clean_text(value: str) -> str:
    return " ".join(value.split()).strip()


def _build_span(
    *,
    source_id: str,
    text: str,
    char_start: int,
    page_number: int | None = None,
    section: str | None = None,
    paragraph_number: int | None = None,
) -> SourceSpan:
    return SourceSpan(
        span_id=_span_id(),
        source_id=source_id,
        text=text,
        page_number=page_number,
        section=section,
        paragraph_number=paragraph_number,
        char_start=char_start,
        char_end=char_start + len(text),
    )


def parse_txt(source_id: str, content: bytes) -> list[SourceSpan]:
    """Parse UTF-8/Windows-compatible text into paragraph-level spans."""

    try:
        decoded = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        try:
            decoded = content.decode("cp1252")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("Text file encoding is not supported.") from exc

    paragraphs = [_clean_text(part) for part in re.split(r"\n\s*\n", decoded)]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    if not paragraphs:
        raise DocumentParseError("The text document does not contain readable text.")

    spans: list[SourceSpan] = []
    cursor = 0
    for number, paragraph in enumerate(paragraphs, start=1):
        spans.append(
            _build_span(
                source_id=source_id,
                text=paragraph,
                page_number=1,
                paragraph_number=number,
                char_start=cursor,
            )
        )
        cursor += len(paragraph) + 2
    return spans


def parse_pdf(source_id: str, content: bytes) -> list[SourceSpan]:
    """Extract one normalized span per readable PDF page."""

    try:
        pdf = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:  # PyMuPDF exposes several parser-specific exceptions.
        raise DocumentParseError("The PDF could not be opened.") from exc

    spans: list[SourceSpan] = []
    cursor = 0
    try:
        for page_index, page in enumerate(pdf, start=1):
            text = _clean_text(page.get_text("text"))
            if not text:
                continue
            spans.append(
                _build_span(
                    source_id=source_id,
                    text=text,
                    page_number=page_index,
                    char_start=cursor,
                )
            )
            cursor += len(text) + 2
    finally:
        pdf.close()

    if not spans:
        raise DocumentParseError(
            "No readable text was found in the PDF. Scanned-PDF OCR is a later adapter."
        )
    return spans


def parse_docx(source_id: str, content: bytes) -> list[SourceSpan]:
    """Extract DOCX paragraphs and table rows with paragraph provenance."""

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("The DOCX file could not be opened.") from exc

    blocks: list[tuple[str, str | None]] = []
    blocks.extend(
        (cleaned, None)
        for paragraph in document.paragraphs
        if (cleaned := _clean_text(paragraph.text))
    )

    for table_index, table in enumerate(document.tables, start=1):
        for row in table.rows:
            row_text = _clean_text(" | ".join(cell.text for cell in row.cells))
            if row_text:
                blocks.append((row_text, f"table_{table_index}"))

    if not blocks:
        raise DocumentParseError("The DOCX document does not contain readable text.")

    spans: list[SourceSpan] = []
    cursor = 0
    for number, (text, section) in enumerate(blocks, start=1):
        spans.append(
            _build_span(
                source_id=source_id,
                text=text,
                section=section,
                paragraph_number=number,
                char_start=cursor,
            )
        )
        cursor += len(text) + 2
    return spans


Parser = Callable[[str, bytes], list[SourceSpan]]

PARSERS: dict[str, Parser] = {
    ".txt": parse_txt,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
}
