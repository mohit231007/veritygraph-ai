import sqlite3
from io import BytesIO

import fitz
import pytest
from app.ingestion.documents import parse_docx, parse_pdf
from app.main import app
from app.repositories.source_repository import SqliteSourceRepository, get_source_repository
from app.services.source_references import (
    extract_docx_hyperlink_references,
    extract_pdf_link_annotation_references,
)
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi.testclient import TestClient

client = TestClient(app)
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TARGET_URL = "https://example.com/research/source#methods"
NORMALIZED_TARGET = "https://example.com/research/source"


@pytest.fixture(autouse=True)
def clear_source_state() -> None:
    get_source_repository().clear()
    yield
    get_source_repository().clear()


def make_docx_with_hidden_hyperlink() -> bytes:
    document = Document()
    paragraph = document.add_paragraph("Read ")
    relationship_id = paragraph.part.relate_to(TARGET_URL, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "the primary report"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    paragraph.add_run(" before reviewing the claim.")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_pdf_with_hidden_uri_annotation() -> bytes:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Read the primary report before reviewing the claim.")
    page.insert_link(
        {
            "kind": fitz.LINK_URI,
            "from": fitz.Rect(72, 58, 230, 82),
            "uri": TARGET_URL,
        }
    )
    content = pdf.tobytes()
    pdf.close()
    return content


def test_docx_hyperlink_relationship_is_retained_without_visible_url() -> None:
    source_id = "src_docx"
    content = make_docx_with_hidden_hyperlink()
    spans = parse_docx(source_id, content)

    assert len(spans) == 1
    assert spans[0].text == "Read the primary report before reviewing the claim."
    assert "https://" not in spans[0].text

    references = extract_docx_hyperlink_references(
        source_id=source_id,
        content=content,
        spans=spans,
    )

    assert len(references) == 1
    reference = references[0]
    assert reference.target_url == TARGET_URL
    assert reference.normalized_target_url == NORMALIZED_TARGET
    assert reference.span_id == spans[0].span_id
    assert reference.paragraph_number == 1
    assert reference.anchor_text == "the primary report"
    assert reference.context_text == spans[0].text
    assert reference.extraction_method == "docx_hyperlink_relationship_v1"


def test_pdf_uri_annotation_is_retained_without_visible_url() -> None:
    source_id = "src_pdf"
    content = make_pdf_with_hidden_uri_annotation()
    spans = parse_pdf(source_id, content)

    assert len(spans) == 1
    assert "https://" not in spans[0].text

    references = extract_pdf_link_annotation_references(
        source_id=source_id,
        content=content,
        spans=spans,
    )

    assert len(references) == 1
    reference = references[0]
    assert reference.target_url == TARGET_URL
    assert reference.normalized_target_url == NORMALIZED_TARGET
    assert reference.span_id == spans[0].span_id
    assert reference.page_number == 1
    assert reference.context_text == spans[0].text
    assert reference.extraction_method == "pdf_link_annotation_v1"


@pytest.mark.parametrize(
    ("filename", "mime_type", "content_factory", "expected_method", "locator_key"),
    [
        (
            "linked.docx",
            DOCX_MIME,
            make_docx_with_hidden_hyperlink,
            "docx_hyperlink_relationship_v1",
            "paragraph_number",
        ),
        (
            "linked.pdf",
            "application/pdf",
            make_pdf_with_hidden_uri_annotation,
            "pdf_link_annotation_v1",
            "page_number",
        ),
    ],
)
def test_document_upload_api_includes_hidden_format_reference(
    filename: str,
    mime_type: str,
    content_factory,
    expected_method: str,
    locator_key: str,
) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": (filename, content_factory(), mime_type)},
    )

    assert response.status_code == 201
    bundle = response.json()
    assert bundle["document"]["metadata"]["reference_count"] == 1
    assert len(bundle["references"]) == 1
    reference = bundle["references"][0]
    assert reference["normalized_target_url"] == NORMALIZED_TARGET
    assert reference["extraction_method"] == expected_method
    assert reference[locator_key] == 1
    assert reference["span_id"] == bundle["spans"][0]["span_id"]


def test_sqlite_migrates_existing_reference_table_with_nullable_locators(tmp_path) -> None:
    database = tmp_path / "legacy.sqlite3"
    with sqlite3.connect(database) as connection:
        connection.execute(
            """
            CREATE TABLE source_references (
                reference_id TEXT PRIMARY KEY,
                source_id TEXT NOT NULL,
                span_id TEXT,
                target_url TEXT NOT NULL,
                normalized_target_url TEXT NOT NULL,
                anchor_text TEXT,
                context_text TEXT,
                extraction_method TEXT NOT NULL
            )
            """
        )

    SqliteSourceRepository(database)

    with sqlite3.connect(database) as connection:
        columns = {
            row[1]
            for row in connection.execute("PRAGMA table_info(source_references)").fetchall()
        }
    assert "page_number" in columns
    assert "paragraph_number" in columns
