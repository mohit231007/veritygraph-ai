from hashlib import sha256
from io import BytesIO

import fitz
import pytest
from app.main import app
from app.repositories.source_repository import get_source_repository
from docx import Document
from fastapi.testclient import TestClient

client = TestClient(app)


@pytest.fixture(autouse=True)
def clear_source_repository() -> None:
    repository = get_source_repository()
    repository.clear()
    yield
    repository.clear()


def test_txt_upload_preserves_source_and_paragraph_provenance() -> None:
    payload = b"Alpha is connected to Beta.\n\nGamma supports Delta."

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("../../research.txt", payload, "text/plain")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["document"]["filename"] == "research.txt"
    assert body["document"]["source_format"] == "txt"
    assert body["document"]["content_hash"] == sha256(payload).hexdigest()
    assert len(body["spans"]) == 2
    assert body["spans"][0]["page_number"] == 1
    assert body["spans"][0]["paragraph_number"] == 1
    assert body["spans"][1]["paragraph_number"] == 2
    assert body["spans"][0]["char_end"] <= body["spans"][1]["char_start"]

    stored = client.get(f"/api/v1/documents/{body['document']['source_id']}")
    assert stored.status_code == 200
    assert stored.json() == body


def test_pdf_upload_preserves_page_number() -> None:
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "VerityGraph evidence lives on page one.")
    payload = pdf.tobytes()
    pdf.close()

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("evidence.pdf", payload, "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["spans"][0]["page_number"] == 1
    assert "VerityGraph evidence lives on page one" in body["spans"][0]["text"]


def test_docx_upload_preserves_paragraphs_and_table_rows() -> None:
    document = Document()
    document.add_paragraph("First evidence paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Entity A"
    table.cell(0, 1).text = "Entity B"
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "brief.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    spans = response.json()["spans"]
    assert spans[0]["text"] == "First evidence paragraph."
    assert spans[1]["section"] == "table_1"
    assert spans[1]["text"] == "Entity A | Entity B"


def test_upload_rejects_unsupported_extension() -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("payload.exe", b"not executable here", "application/octet-stream")},
    )

    assert response.status_code == 415
    assert "PDF, DOCX or TXT" in response.json()["detail"]


def test_upload_rejects_empty_document() -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("empty.txt", b"", "text/plain")},
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "The uploaded document is empty."


def test_missing_source_returns_404() -> None:
    response = client.get("/api/v1/documents/src_missing")

    assert response.status_code == 404
    assert response.json()["detail"] == "Source not found."
